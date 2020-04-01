# 读取目录下所有的*.cpp，*.h，*.rc文件，并写入word中。（存在问题：无法处理ValueError问题，需要后续确认） 
import os
import os.path

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

def readFile(filepath, document):
    list_name = os.listdir(filepath)
    print(list_name)
    list_dir = []
    for name in list_name:
        if os.path.isdir(filepath + "/" + name):
            list_dir.append(name)
        elif os.path.isfile(filepath + "/" + name) \
                and (os.path.splitext(filepath + "/" + name)[-1] == ".cpp"
                     or os.path.splitext(filepath + "/" + name)[-1] == ".h"
                     or os.path.splitext(filepath + "/" + name)[-1] == ".rc"):
            try:
                print(filepath + "/" + name)
                f = open(filepath + "/" + name, 'r')
                str_content = f.read()
                f.close()
                document.add_heading(filepath + "/" + name, level=3)
                document.add_paragraph(str_content)
            except ValueError:
                print(filepath + "/" + name, " is Vaule error!!!")
                f.close()
            except UnicodeDecodeError:
                try:
                    print(filepath + "/" + name, " tried utf-8!")
                    f = open(filepath + "/" + name, 'r', encoding="utf-8")
                    str_content = f.read()
                    f.close()
                    document.add_heading(filepath + "/" + name, level=3)
                    document.add_paragraph(str_content)
                except UnicodeDecodeError:
                    try:
                        print(filepath + "/" + name, " tried GB2312!")
                        f = open(filepath + "/" + name, 'r', encoding="GB2312")
                        str_content = f.read()
                        f.close()
                        document.add_heading(filepath + "/" + name, level=3)
                        document.add_paragraph(str_content)
                    except UnicodeDecodeError:
                        print(filepath + "/" + name, " is error!!!")
                        f.close()

    for dir_item in list_dir:
        readFile(filepath + "/" + dir_item, document)
    #             document.add_paragraph(filepath + "/" + dir_item)
    #             document.add_heading(dir_item, level=3)

    document.save("./test_word.docx")



if __name__ == "__main__":
    document = Document()
    readFile("./Client", document)